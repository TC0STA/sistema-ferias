from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "documentacao"
OUTPUT = OUTPUT_DIR / "Especificacao_Funcional_Modulo_Importacao_v1.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B3A66"
INK = "172033"
MUTED = "5D6878"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D2DE"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name="Calibri",
    size=11,
    color=INK,
    bold=False,
    italic=False
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr

    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER, size="6"):
    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = 91
    num_id = 91

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    level.extend((start, num_fmt, level_text, justification, p_pr))
    abstract_num.append(level)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_element))


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_body(doc, text, *, bold_prefix=None, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.line_spacing = 1.1
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge, size, color in (
        ("top", "6", "AFC8E5"),
        ("left", "18", BLUE),
        ("bottom", "6", "AFC8E5"),
        ("right", "6", "AFC8E5"),
    ):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), color)
        borders.append(border)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color=INK)


def add_rule(doc, code, title, requirement, acceptance):
    heading = doc.add_paragraph(style="Heading 2")
    code_run = heading.add_run(f"{code}  ")
    set_run_font(code_run, size=13, color=DARK_BLUE, bold=True)
    title_run = heading.add_run(title)
    set_run_font(title_run, size=13, color=BLUE, bold=True)
    add_body(doc, requirement)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_after = Pt(8)
    label = paragraph.add_run("Critério de aceite: ")
    set_run_font(label, size=10.5, color=MUTED, bold=True)
    text = paragraph.add_run(acceptance)
    set_run_font(text, size=10.5, color=MUTED)


def add_log_table(doc):
    rows = [
        ("Arquivo", "Nome original do arquivo enviado pelo RH."),
        ("Data", "Data em que a tentativa de importação ocorreu."),
        ("Hora", "Horário da tentativa de importação."),
        ("Usuário", "Responsável autenticado pela operação."),
        ("Quantidade de registros", "Total de registros identificados na planilha."),
        ("Tempo de processamento", "Duração total do processamento."),
        ("Resultado", "Sucesso, falha ou cancelamento."),
        ("Quantidade de erros", "Total de erros encontrados na validação."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    row_pr.append(header_marker)
    header = table.rows[0].cells
    for index, text in enumerate(("Campo", "Descrição")):
        set_cell_shading(header[index], LIGHT_GRAY)
        paragraph = header[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    for field, description in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 2700)
        set_cell_width(cells[1], 6660)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        first = cells[0].paragraphs[0]
        first.paragraph_format.space_after = Pt(0)
        set_run_font(first.add_run(field), size=10.5, bold=True)
        second = cells[1].paragraphs[0]
        second.paragraph_format.space_after = Pt(0)
        set_run_font(second.add_run(description), size=10.5)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    number_id = add_numbering_definition(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(
        header.add_run("FOKUS FÉRIAS  |  ESPECIFICAÇÃO FUNCIONAL"),
        size=9,
        color=MUTED,
        bold=True
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("Documento EF-MI-001  |  Página "), size=9, color=MUTED)
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(
        kicker.add_run("DOCUMENTO OFICIAL DO PROJETO"),
        size=10,
        color=BLUE,
        bold=True
    )

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    set_run_font(
        title.add_run("Especificação Funcional"),
        size=24,
        color=DARK_BLUE,
        bold=True
    )

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(
        subtitle.add_run("Módulo de Importação — Fokus Férias"),
        size=14,
        color=MUTED
    )

    metadata = (
        ("Código", "EF-MI-001"),
        ("Versão", "1.0"),
        ("Data", "29/07/2026"),
        ("Status", "Documento oficial"),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_heading("1. Objetivo", level=1)
    add_body(
        doc,
        "Receber a planilha enviada pelo RH, validar todas as informações e "
        "disponibilizar os dados para que a TI execute o processo de bloqueio "
        "dos usuários de forma segura e organizada."
    )
    add_callout(
        doc,
        "Princípio fundamental",
        "O fluxo principal definido neste documento é imutável e constitui o "
        "coração operacional do sistema Fokus Férias."
    )

    doc.add_heading("2. Escopo funcional", level=1)
    add_body(
        doc,
        "Esta especificação abrange o recebimento, a validação, a importação, "
        "a proteção dos dados, a atualização dos módulos consumidores e o "
        "registro integral da operação."
    )

    doc.add_heading("3. Fluxo principal", level=1)
    stages = [
        ("RH envia a planilha", "O arquivo é disponibilizado para processamento."),
        ("Selecionar arquivo", "O usuário escolhe o arquivo Excel recebido."),
        ("Validar", "O sistema verifica estrutura, campos e conteúdo."),
        ("Importar", "A operação é autorizada somente após validação bem-sucedida."),
        ("Atualizar banco", "Os dados validados são persistidos com segurança."),
        ("Atualizar Dashboard", "Os indicadores passam a refletir a nova base."),
        ("Atualizar Histórico", "A importação e sua versão ficam registradas."),
        ("Atualizar Auditoria", "A tentativa e seu resultado são rastreados."),
        (
            "Disponibilizar usuários para bloqueio",
            "A TI recebe a relação operacional resultante da importação."
        ),
    ]
    for stage, detail in stages:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, number_id)
        paragraph.paragraph_format.keep_together = True
        set_run_font(paragraph.add_run(stage), bold=True, color=DARK_BLUE)
        set_run_font(paragraph.add_run(f" — {detail}"))

    doc.add_page_break()
    doc.add_heading("4. Regras de negócio", level=1)
    add_rule(
        doc,
        "RN001",
        "Formato do arquivo",
        "A importação deve aceitar exclusivamente arquivos Excel.",
        "Arquivos fora dos formatos Excel permitidos são recusados antes da leitura."
    )
    add_rule(
        doc,
        "RN002",
        "Validação estrutural",
        "A estrutura da planilha deve ser validada antes do processamento dos registros.",
        "Nenhum registro é processado enquanto a estrutura não for aprovada."
    )
    add_rule(
        doc,
        "RN003",
        "Campos obrigatórios",
        "A planilha deve conter, no mínimo, os campos Nome, Data Início e Data Fim. "
        "A relação será adaptada ao layout oficial fornecido pelo RH.",
        "A ausência de qualquer campo obrigatório produz erro estrutural."
    )
    add_rule(
        doc,
        "RN004",
        "Importação atômica",
        "Se existir erro estrutural, a importação deve ser cancelada integralmente. "
        "Não é permitida importação parcial da planilha.",
        "Após uma reprovação estrutural, nenhum dado do arquivo é gravado."
    )
    add_rule(
        doc,
        "RN005",
        "Auditoria obrigatória",
        "Toda tentativa de importação deve gerar um registro de auditoria, inclusive "
        "quando houver erro ou cancelamento.",
        "Cada tentativa possui um evento de auditoria com usuário, data, hora e resultado."
    )
    add_rule(
        doc,
        "RN006",
        "Backup obrigatório",
        "Antes de qualquer gravação no banco de dados, o sistema deve criar um backup.",
        "Se o backup não puder ser criado, a importação é interrompida sem alterar a base."
    )
    add_rule(
        doc,
        "RN007",
        "Atualização automática",
        "Após uma importação concluída, o sistema deve atualizar automaticamente "
        "Dashboard, Histórico, Relatórios e Centro de Operações.",
        "Os quatro módulos refletem a nova importação sem intervenção manual."
    )

    doc.add_page_break()
    doc.add_heading("5. Estados da importação", level=1)
    add_body(
        doc,
        "Cada tentativa deve possuir um estado identificável durante todo o processamento."
    )
    doc.add_heading("5.1 Fluxo concluído", level=2)
    for state, description in (
        ("Recebido", "Arquivo recebido pelo sistema."),
        ("Validando", "Estrutura e conteúdo em análise."),
        ("Pronto para importar", "Validação concluída sem impedimentos."),
        ("Importando", "Backup realizado e dados em processamento."),
        ("Concluído", "Importação finalizada e módulos atualizados."),
    ):
        add_body(doc, f"{state}: {description}", bold_prefix=f"{state}: ", after=4)

    doc.add_heading("5.2 Fluxo cancelado", level=2)
    for state, description in (
        ("Recebido", "Arquivo recebido pelo sistema."),
        ("Validando", "Estrutura e conteúdo em análise."),
        ("Erro encontrado", "Um ou mais impedimentos foram identificados."),
        ("Importação cancelada", "Nenhum dado do arquivo foi gravado."),
    ):
        add_body(doc, f"{state}: {description}", bold_prefix=f"{state}: ", after=4)

    add_callout(
        doc,
        "Garantia de integridade",
        "A transição para “Importando” somente pode ocorrer depois da validação "
        "bem-sucedida e da criação do backup obrigatório."
    )

    doc.add_heading("6. Log da importação", level=1)
    add_body(
        doc,
        "Toda tentativa de importação deve registrar os seguintes dados para "
        "auditoria, suporte e rastreabilidade:"
    )
    add_log_table(doc)

    doc.add_heading("7. Critérios gerais de aceite", level=1)
    criteria = [
        "Arquivos não Excel são rejeitados sem alteração da base.",
        "Erros estruturais impedem integralmente a importação.",
        "Toda tentativa, com sucesso ou falha, produz auditoria.",
        "Nenhuma gravação ocorre sem backup prévio confirmado.",
        "Importações concluídas atualizam os módulos consumidores automaticamente.",
        "O log contém todos os campos definidos na seção 6.",
        "O estado final da tentativa é sempre Concluído ou Importação cancelada.",
    ]
    for criterion in criteria:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        set_run_font(paragraph.add_run("☐  "), color=BLUE, bold=True)
        set_run_font(paragraph.add_run(criterion))

    doc.add_heading("8. Controle do documento", level=1)
    add_body(
        doc,
        "Este documento é a referência funcional oficial do Módulo de Importação. "
        "Qualquer alteração nas regras, nos estados ou no fluxo principal deve ser "
        "registrada em uma nova versão desta especificação."
    )
    add_callout(
        doc,
        "Regra de governança",
        "O fluxo principal não deve ser alterado silenciosamente por implementação, "
        "ajuste de interface ou nova funcionalidade."
    )

    core = doc.core_properties
    core.title = "Especificação Funcional – Módulo de Importação"
    core.subject = "Documento oficial do projeto Fokus Férias"
    core.author = "Projeto Fokus Férias"
    core.keywords = "Fokus Férias, importação, especificação funcional"
    core.comments = "Versão 1.0 — 29/07/2026"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
